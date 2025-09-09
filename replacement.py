from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re
    
def replace_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:

        style = paragraph.style
        runs = paragraph.runs
        
        full_text = ''.join(run.text for run in runs)
        print(full_text)

        if old_text in full_text:
            paragraph.clear()
            parts = full_text.split(old_text)
            
            for i, part in enumerate(parts):
                if part:  
                    paragraph.add_run(part)
                
                if i < len(parts) - 1:
                    run = paragraph.add_run(new_text)
                    if runs:
                        first_run = runs[0]
                        run.bold = first_run.bold
                        run.italic = first_run.italic
                        run.underline = first_run.underline
                        if first_run.font.color:
                            run.font.color.rgb = first_run.font.color.rgb
                        if first_run.font.name:
                            run.font.name = first_run.font.name
                        if first_run.font.size:
                            run.font.size = first_run.font.size
            paragraph.style = style

def replace_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, old_text, new_text)
            for nested_table in cell.tables:
                replace_in_table(nested_table, old_text, new_text)


def replace_text_in_docx(doc, old_text, new_text):

    # doc = Document(input_file)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, old_text, new_text)
    
    for table in doc.tables:
        replace_in_table(table, old_text, new_text)
    
    
    # doc.save(output_file)


def replace(docx, dict_replacements):
    for key in dict_replacements:
        replace_text_in_docx(docx, key, dict_replacements[key])

if __name__ == "__main__":
    doc = Document("Реферат.docx")
    replace_text_in_docx(doc, "__authors__", "фувертпвот")
    doc.save("output_file.docx")
    


# Для чего нужен модуль. Модуль используется для...
# Что использует модуль для работы 
# Область применения данного модуля

# __authors__
# __director__
# __program__
# __annotation__
# __type__
# __language__
# __os__
# __memory__

